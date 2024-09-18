import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import docx.opc.constants
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
import win32com.client
import logging

# Set up logging configuration
logging.basicConfig(level=logging.INFO, format='%(message)s')

def add_hyperlink(run, url, text):
    """
    Adds a hyperlink to a run in a Word document.
    """
    run.font.underline = True
    r_id = run.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    new_run.append(r_pr)
    new_run.text = text
    hyperlink.append(new_run)
    run._r.append(hyperlink)

def read_excel_file(file_path):
    """
    Reads the Excel file containing Jira issues and returns a list of dictionaries.
    """
    try:
        df = pd.read_excel(file_path)
        df = df.fillna('').astype(str)
        data = df.to_dict('records')
        return data
    except Exception as e:
        logging.error(f"Error reading Excel file '{file_path}': {e}")
        return []

def process_data(data):
    """
    Processes the data from Excel, organizing it into a hierarchical structure.
    """
    structured_data = {}
    current_theme = None
    current_goal = None
    current_status = None
    current_initiative = None

    for row in data:
        issue_type = row.get('Issue Type', '')
        issue_key = row.get('Key', '')
        summary = row.get('Summary', '')
        hebrew_summary = row.get('Hebrew Summary', '')
        status = row.get('Status', '')
        description = row.get('Description', '')
        start_date = row.get('Start date', '')
        due_date = row.get('Due date', '')

        if issue_type == 'Theme':
            # Start a new theme
            current_theme = issue_key
            structured_data[current_theme] = {
                'summary': summary,
                'hebrew_summary': hebrew_summary,
                'goals': {}
            }
            current_goal = None
            current_status = None
            current_initiative = None
        elif issue_type == 'Goal':
            # Add a goal under the current theme
            if current_theme:
                current_goal = issue_key
                structured_data[current_theme]['goals'][current_goal] = {
                    'summary': summary,
                    'hebrew_summary': hebrew_summary,
                    'description': description,
                    'statuses': {}
                }
                current_status = None
                current_initiative = None
            else:
                logging.warning(f"Goal '{issue_key}' found without a current theme.")
        elif issue_type == 'Initiative':
            # Add an initiative under the current goal
            if current_theme and current_goal:
                current_status = status
                current_initiative = issue_key
                statuses = structured_data[current_theme]['goals'][current_goal]['statuses']
                if current_status not in statuses:
                    statuses[current_status] = {}
                statuses[current_status][current_initiative] = {
                    'summary': summary,
                    'hebrew_summary': hebrew_summary,
                    'description': description,
                    'start_date': start_date,
                    'due_date': due_date,
                    'leads': {}
                }
            else:
                logging.warning(f"Initiative '{issue_key}' found without a current theme and goal.")
        elif issue_type == 'Lead':
            # Add a lead under the current initiative
            if current_theme and current_goal and current_status and current_initiative:
                leads = structured_data[current_theme]['goals'][current_goal]['statuses'][current_status][current_initiative]['leads']
                leads[issue_key] = {
                    'summary': summary,
                    'hebrew_summary': hebrew_summary,
                    'description': description
                }
            else:
                logging.warning(f"Lead '{issue_key}' found without a current theme, goal, status, and initiative.")
        elif issue_type == '':
            # Stop processing if a row with "Not an issue" is found
            if summary == "Not an issue":
                break
        else:
            logging.warning(f"Unknown issue type '{issue_type}' for key '{issue_key}'.")

    return structured_data

def create_word_document(structured_data, output_file_path, include_todo=False):
    """
    Creates a Word document from the structured data.
    """
    doc = Document()
    setup_document(doc)
    add_content(doc, structured_data, include_todo)
    success = save_document(doc, output_file_path)
    return success

def setup_document(doc):
    """
    Sets up the document layout and styles.
    """
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    styles = doc.styles
    heading_styles = [
        ('Heading 1', 22, 0),
        ('Heading 2', 20, 1),
        ('Heading 3', 18, 0)
    ]
    for style_name, size, indent in heading_styles:
        style = styles[style_name]
        style.font.size = Pt(size)
        style.paragraph_format.left_indent = Cm(indent)

    add_toc_field(doc)

def add_toc_field(doc):
    """
    Adds a Table of Contents field to the document.
    """
    doc.add_heading("Roadmap Status Report", level=0)
    doc.add_paragraph("Table of Contents", style='TOC Heading')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, fld_char_end])

    reminder = doc.add_paragraph()
    reminder_run = reminder.add_run("Right-click and select 'Update Field' to update the Table of Contents.")
    reminder_run.font.italic = True
    reminder_run.font.size = Pt(9)
    reminder_run.font.color.rgb = RGBColor(128, 128, 128)

def add_content(doc, structured_data, include_todo):
    """
    Adds the content to the document based on the structured data.
    """
    status_order = ['Done', 'In Progress', 'Next']
    if include_todo:
        status_order.append('To Do')

    for theme_key, theme_data in structured_data.items():
        theme_printed = False
        for goal_key, goal_data in theme_data['goals'].items():
            goal_printed = False
            for status in status_order:
                if status in goal_data['statuses'] and goal_data['statuses'][status]:
                    theme_printed, goal_printed = add_theme_goal_content(
                        doc, theme_key, theme_data, goal_key, goal_data, status, theme_printed, goal_printed
                    )

def add_theme_goal_content(doc, theme_key, theme_data, goal_key, goal_data, status, theme_printed, goal_printed):
    if not theme_printed:
        heading = doc.add_heading(level=1)
        heading.add_run(f"{theme_data['summary']} (")
        add_hyperlink(heading.add_run(), f"https://omnisys.atlassian.net/browse/{theme_key}", theme_key)
        heading.add_run(")")

        # Add Hebrew summary with RTL control characters
        hebrew_text = '\u202B' + theme_data['hebrew_summary'] + '\u202C'
        hebrew_summary = doc.add_paragraph()
        hebrew_run = hebrew_summary.add_run(hebrew_text)
        hebrew_run.font.size = Pt(12)
        hebrew_summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hebrew_summary.paragraph_format.bidi = True
        theme_printed = True

    if not goal_printed:
        heading = doc.add_heading(level=2)
        heading.add_run(f"{goal_data['summary']} (")
        add_hyperlink(heading.add_run(), f"https://omnisys.atlassian.net/browse/{goal_key}", goal_key)
        heading.add_run(")")

        # Add Hebrew summary with RTL control characters
        hebrew_text = '\u202B' + goal_data['hebrew_summary'] + '\u202C'
        hebrew_summary = doc.add_paragraph()
        hebrew_run = hebrew_summary.add_run(hebrew_text)
        hebrew_run.font.size = Pt(12)
        hebrew_summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hebrew_summary.paragraph_format.bidi = True
        goal_printed = True

    add_status_table(doc, status, goal_data['statuses'][status])
    return theme_printed, goal_printed

def add_status_table(doc, status, initiatives):
    """
    Adds a table for the status and its initiatives to the document.
    """
    doc.add_heading(f"Status: {status}", level=3)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Title'
    hdr_cells[1].text = 'Description'

    for initiative_key, initiative_data in initiatives.items():
        row_cells = table.add_row().cells
        add_initiative_to_table(row_cells, initiative_key, initiative_data)

    doc.add_paragraph()

def add_initiative_to_table(row_cells, initiative_key, initiative_data):
    for cell in row_cells:
        cell.paragraphs[0].style.font.size = Pt(12)

    # Add initiative summary and key
    summary_paragraph = row_cells[0].paragraphs[0]
    summary_paragraph.clear()
    summary_paragraph.style.font.size = Pt(12)
    summary_paragraph.add_run(f"{initiative_data['summary']} (")
    add_hyperlink(summary_paragraph.add_run(), f"https://omnisys.atlassian.net/browse/{initiative_key}", initiative_key)
    summary_paragraph.add_run(")")

    # Add Hebrew summary with RTL control characters
    hebrew_text = '\u202B' + initiative_data['hebrew_summary'] + '\u202C'
    hebrew_summary = row_cells[0].add_paragraph()
    hebrew_run = hebrew_summary.add_run(hebrew_text)
    hebrew_run.font.size = Pt(12)
    hebrew_summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hebrew_summary.paragraph_format.bidi = True

    # Add start date and due date
    start_date = initiative_data['start_date'].split()[0] if initiative_data['start_date'] else 'Unknown'
    due_date = initiative_data['due_date'].split()[0] if initiative_data['due_date'] else 'Unknown'
    dates_paragraph = row_cells[0].add_paragraph()
    dates_paragraph.add_run(f"Start Date: {start_date}\nDue Date: {due_date}")
    dates_paragraph.style.font.size = Pt(12)

    # Add description
    description_paragraph = row_cells[1].paragraphs[0]
    description_paragraph.text = initiative_data['description']
    description_paragraph.style.font.size = Pt(12)

    # Add linked initiatives
    if initiative_data['leads']:
        linked_initiatives = row_cells[1].add_paragraph("\n\nLinked initiatives:")
        linked_initiatives.style.font.size = Pt(12)
        for lead_key, lead_data in initiative_data['leads'].items():
            p = row_cells[1].add_paragraph("- ")
            p.add_run(f"{lead_data['summary']} (")
            add_hyperlink(p.add_run(), f"https://omnisys.atlassian.net/browse/{lead_key}", lead_key)
            p.add_run(")")
            p.style.font.size = Pt(12)

def save_document(doc, output_file_path):
    """
    Saves the Word document to the specified file path.
    """
    try:
        doc.save(output_file_path)
    except PermissionError:
        logging.error(f"PermissionError: Unable to save '{output_file_path}'. Please close the file if it's open and try again.")
        return False
    except Exception as e:
        logging.error(f"Error saving document: {e}")
        return False
    return True

def update_toc(doc_path):
    """
    Updates the Table of Contents in the Word document.
    """
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        doc.TablesOfContents(1).Update()
        doc.Save()
        doc.Close()
        word.Quit()
        logging.info("Table of Contents updated successfully.")
    except Exception as e:
        logging.error(f"Error updating Table of Contents: {e}")

def main():
    """
    Main function to generate the roadmap report.
    """
    excel_files = [f for f in os.listdir() if re.match(r"Roadmap.*\.xlsx", f)]
    if not excel_files:
        logging.error("No matching Excel file found.")
        return

    file_path = excel_files[0]
    data = read_excel_file(file_path)
    if not data:
        logging.error("No data read from the Excel file.")
        return

    logging.info(f"Successfully read {len(data)} rows from '{file_path}'")
    include_todo = False  # Set this to True if you want to include 'To Do' status
    structured_data = process_data(data)
    output_file_name = "Roadmap Status Report.docx"
    output_file_path = os.path.abspath(output_file_name)
    success = create_word_document(structured_data, output_file_path, include_todo)
    if success:
        logging.info(f"\nWord document created: '{output_file_path}'")
        update_toc(output_file_path)
    else:
        logging.error("\nFailed to create Word document. Please check the error messages above.")

if __name__ == "__main__":
    main()
